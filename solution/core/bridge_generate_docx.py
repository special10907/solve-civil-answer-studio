#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Bridge Script: Node.js -> Professional DOCX Generation
- Takes a JSON input containing structured question/answer data.
- Uses subnote_framework.py to generate a professional .docx subnote.
"""

import sys
import json
import re
import traceback
import base64
import binascii
import struct
import zlib
from pathlib import Path
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from docx.shared import Cm
import requests
# MCP LLM API 호출 함수
def llm_api_call(prompt):
    mcp_url = "https://edge.flowith.io/external/use/llm"
    payload = {
        "prompt": prompt,
        "model": "gpt-4o",
        "max_tokens": 2048
    }
    try:
        response = requests.post(mcp_url, json=payload)
        if response.status_code == 200:
            return response.json()
        else:
            print(f"MCP LLM API 호출 실패: {response.status_code} {response.text}")
            return None
    except Exception as e:
        print(f"MCP LLM API 호출 오류: {e}")
        return None

# Add research directory to sys.path
ROOT_DIR = Path(__file__).resolve().parent
WORKSPACE_ROOT = ROOT_DIR.parent.parent
RESEARCH_DIR = ROOT_DIR / "research"
sys.path.insert(0, str(RESEARCH_DIR))

try:
    from subnote_framework import (
        create_document, save_document,
        add_exam_header, add_section_header, add_footer,
        add_content_line, add_content_sub,
        add_data_table, add_insight_box, add_divider, add_strategy_section,
        add_callout_box, TABLE_W
    )
except ImportError as e:
    print(
        json.dumps(
            {"ok": False, "error": f"Import error: {e}"},
            ensure_ascii=False,
        )
    )
    sys.exit(1)


def _ensure_diagram_content_rule(content):
    text = str(content or "").strip()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    base = "\n".join(lines)

    required = [
        ("도해 목적", "핵심 메커니즘을 빠르게 전달"),
        ("구성 요소", "하중·저항·경계조건·핵심 부재"),
        ("작성 순서", "1) 외곽 2) 하중/반력 3) 내부 경로/레이블"),
        ("채점 포인트", "기준 용어 일치, 수치/근거, 결론 연결"),
    ]

    out = list(lines)
    for label, fallback in required:
        if not re.search(rf"{re.escape(label)}\s*:", base, re.IGNORECASE):
            out.append(f"{label}: {fallback}")

    return "\n".join(out).strip()


def _normalize_diagrams(llm_data, title):
    safe_title = str(title or "핵심 구조 메커니즘").strip() or "핵심 구조 메커니즘"
    incoming = (
        llm_data.get("diagrams", []) if isinstance(llm_data, dict) else []
    )
    incoming = incoming if isinstance(incoming, list) else []

    defaults = [
        {
            "title": f"{safe_title} 메커니즘 도해",
            "content": (
                "도해 목적: 문제의 하중-저항 흐름을 시각화\n"
                "구성 요소: 작용하중, 지점반력, 주요 부재, 위험 구간\n"
                "작성 순서: 1) 형상/경계조건 2) 하중·반력 3) 응력·힘 흐름\n"
                "채점 포인트: 핵심 용어 표기, 방향성, 결론과의 연결"
            ),
        },
        {
            "title": f"{safe_title} 대안 비교표",
            "content": (
                "도해 목적: 대안별 설계 판단근거를 표로 비교\n"
                "구성 요소: 비교항목(안전성·시공성·경제성·유지관리성), 대안 A/B\n"
                "작성 순서: 1) 비교항목 설정 2) 대안별 평가 3) 권고안 도출\n"
                "채점 포인트: 기준 근거, 트레이드오프 설명, 실무 제언"
            ),
        },
    ]

    normalized = []
    for idx, item in enumerate(incoming):
        if not isinstance(item, dict):
            continue
        item_title = str(item.get("title") or "").strip() or f"도해 {idx + 1}"
        item_content = _ensure_diagram_content_rule(item.get("content") or "")
        if not item_content:
            continue
        normalized.append({"title": item_title, "content": item_content})

    while len(normalized) < 2:
        seed = defaults[len(normalized)]
        normalized.append(
            {
                "title": seed["title"],
                "content": _ensure_diagram_content_rule(seed["content"]),
            }
        )

    return normalized[:4]


def _is_noise_line(text):
    line = str(text or "").strip()
    if not line:
        return True
    noise_patterns = [
        r"^\[\s*심화\s*보강",
        r"^\[\s*web_research\s*\]$",
        r"^\[\s*mandatory_pipeline_context\s*\]$",
        r"^\[\s*deep_research_parsed\s*\]$",
        r"^\[\s*검색\s*컨텍스트\s*요약\s*\]$",
        r"^query\s*:",
        r"^status\s*:",
        r"^message\s*:",
        r"^title\s*:",
        r"^summary\s*:",
        r"^url\s*:",
        r"^references\s*:",
        r"^참고\s*링크\s*없음$",
        r"^[-*]\s*요청사항\s*:",
        r"^[-*]\s*탐색소스\s*:",
        r"^근거첨부$",
        r"^\|\s*단계\s*\|\s*소스\s*\|",
        r"^\|\s*---",
        r"^\|\s*[1-5]\s*\|",
        r"시각화\s*요약",
        r"^\d+\)\s*(그림|표|그래프)\s*[:：]",
        r"^[-*•]\s*(목적|작성\s*기준|채점\s*포인트)\s*[:：]",
        r"^\[?OCR\]?",
        r"텍스트\s*레이어\s*(없음|미존재|누락)",
        r"image\s*text\s*layer\s*(missing|not\s*found)",
    ]
    for pattern in noise_patterns:
        if re.search(pattern, line, re.IGNORECASE):
            return True
    return False


def _clean_multiline_text(text, max_lines=22):
    raw = str(text or "").replace("\r", "\n")
    lines = [ln.strip() for ln in raw.split("\n")]
    cleaned = [ln for ln in lines if ln and not _is_noise_line(ln)]
    return cleaned[:max_lines]


def _clean_inline_text(text, fallback=""):
    lines = _clean_multiline_text(text, max_lines=2)
    if not lines:
        return str(fallback or "").strip()
    return " ".join(lines).strip()


def _normalize_visual_items(llm_data, title):
    safe_title = str(title or "핵심 구조 메커니즘").strip() or "핵심 구조 메커니즘"
    visuals = []

    if isinstance(llm_data, dict):
        raw_visuals = llm_data.get("visuals", [])
        if isinstance(raw_visuals, list):
            for item in raw_visuals:
                if not isinstance(item, dict):
                    continue
                kind = str(item.get("kind") or "diagram").strip().lower()
                visuals.append(
                    {
                        "kind": kind,
                        "title": (
                            str(item.get("title") or "").strip()
                            or f"{safe_title} 시각자료"
                        ),
                        "purpose": str(item.get("purpose") or "").strip(),
                        "spec": str(item.get("spec") or "").strip(),
                        "scoringPoint": (
                            str(item.get("scoringPoint") or "").strip()
                        ),
                        "imageData": str(item.get("imageData") or "").strip(),
                        "imageUrl": str(item.get("imageUrl") or "").strip(),
                    }
                )

    if not visuals:
        diagrams = _normalize_diagrams(llm_data, title)
        for d in diagrams:
            visuals.append(
                {
                    "kind": "diagram",
                    "title": str(d.get("title") or "도해").strip(),
                    "purpose": "핵심 메커니즘 전달",
                    "spec": str(d.get("content") or "").strip(),
                    "scoringPoint": "용어 일치·근거 연결",
                    "imageData": "",
                    "imageUrl": "",
                }
            )

    normalized = []
    for item in visuals:
        kind = item.get("kind", "diagram")
        if kind not in {"diagram", "table", "graph", "image"}:
            kind = "diagram"
        normalized.append({**item, "kind": kind})

    return normalized[:6]


def _render_visual_table(doc, item, is_submission=False, caption=""):
    title = _clean_inline_text(item.get("title"), "비교표")
    spec_text = item.get("spec") or ""
    purpose = _clean_inline_text(item.get("purpose"))
    scoring = _clean_inline_text(item.get("scoringPoint"))

    label = caption or "표"
    if is_submission:
        add_content_sub(doc, f"{label}: {title}", indent=0.8)
    else:
        add_content_line(doc, 1, "", bold_prefix=f"[{label}] {title}: ")
    if purpose:
        add_content_sub(doc, f"목적: {purpose}", indent=1.0)

    rows = []
    for line in _clean_multiline_text(spec_text, max_lines=12):
        if re.match(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?$", line):
            continue
        if "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if len(cells) >= 2:
                rows.append(cells)
        elif ":" in line:
            left, right = [p.strip() for p in line.split(":", 1)]
            if left and right:
                rows.append([left, right])

    if len(rows) < 2:
        rows = [
            ["비교항목", "대안 A", "대안 B"],
            ["안전성", "기준 만족", "보강 필요"],
            ["시공성", "보통", "우수"],
            ["경제성", "초기비용 낮음", "유지관리비 낮음"],
        ]

    headers = rows[0]
    data = rows[1:] if len(rows) > 1 else [["항목", "내용"]]
    col_count = max(len(headers), max((len(r) for r in data), default=2))

    headers = headers + [f"열{idx+1}" for idx in range(len(headers), col_count)]
    norm_data = []
    for row in data:
        norm_data.append(row + ["-"] * (col_count - len(row)))

    if col_count == 2:
        col_widths = [4.0, TABLE_W - 4.0]
    elif col_count == 3:
        col_widths = [3.5, 5.0, TABLE_W - 8.5]
    else:
        each = TABLE_W / float(col_count)
        col_widths = [each for _ in range(col_count)]

    add_data_table(
        doc,
        headers=headers,
        data=norm_data[:8],
        col_widths=col_widths,
    )
    if scoring:
        add_content_sub(doc, f"채점 포인트: {scoring}", indent=1.0)


def _write_png(path, width, height, painter):
    pixels = bytearray([255] * (width * height * 3))

    def set_px(x, y, r, g, b):
        if 0 <= x < width and 0 <= y < height:
            idx = (y * width + x) * 3
            pixels[idx: idx + 3] = bytes((r, g, b))

    painter(set_px)

    raw = bytearray()
    stride = width * 3
    for y in range(height):
        raw.append(0)
        start = y * stride
        raw.extend(pixels[start:start + stride])

    def chunk(tag, data):
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    png = bytearray(b"\x89PNG\r\n\x1a\n")
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    png.extend(chunk(b"IHDR", ihdr))
    png.extend(chunk(b"IDAT", zlib.compress(bytes(raw), 9)))
    png.extend(chunk(b"IEND", b""))

    with open(path, "wb") as f:
        f.write(png)


def _generate_placeholder_visual_png(path, kind="diagram"):
    width, height = 1080, 520

    def painter(set_px):
        def rect(x1, y1, x2, y2, color):
            r, g, b = color
            for yy in range(max(0, y1), min(height, y2)):
                for xx in range(max(0, x1), min(width, x2)):
                    set_px(xx, yy, r, g, b)

        def line(x1, y1, x2, y2, color):
            r, g, b = color
            dx = abs(x2 - x1)
            dy = -abs(y2 - y1)
            sx = 1 if x1 < x2 else -1
            sy = 1 if y1 < y2 else -1
            err = dx + dy
            x, y = x1, y1
            while True:
                for ox in (-1, 0, 1):
                    for oy in (-1, 0, 1):
                        set_px(x + ox, y + oy, r, g, b)
                if x == x2 and y == y2:
                    break
                e2 = 2 * err
                if e2 >= dy:
                    err += dy
                    x += sx
                if e2 <= dx:
                    err += dx
                    y += sy

        rect(0, 0, width, height, (248, 250, 252))
        rect(40, 40, width - 40, height - 40, (255, 255, 255))

        if kind == "graph":
            line(140, 420, 940, 420, (70, 70, 70))
            line(140, 420, 140, 120, (70, 70, 70))
            line(160, 390, 320, 330, (37, 99, 235))
            line(320, 330, 520, 300, (37, 99, 235))
            line(520, 300, 760, 220, (37, 99, 235))
            line(760, 220, 900, 180, (37, 99, 235))
        else:
            rect(130, 170, 360, 300, (219, 234, 254))
            rect(440, 170, 670, 300, (254, 226, 226))
            rect(750, 170, 980, 300, (220, 252, 231))
            line(360, 235, 440, 235, (51, 65, 85))
            line(670, 235, 750, 235, (51, 65, 85))

    _write_png(path, width, height, painter)


def _decode_image_data_to_file(image_data, out_path):
    src = str(image_data or "").strip()
    if not src:
        return False

    # data URL 허용: data:image/png;base64,xxxx
    if "," in src and src.lower().startswith("data:image"):
        src = src.split(",", 1)[1].strip()

    try:
        raw = base64.b64decode(src, validate=True)
    except (binascii.Error, ValueError):
        return False

    if not raw:
        return False

    with open(out_path, "wb") as f:
        f.write(raw)
    return True


def _download_image_from_url(image_url, out_path):
    url = str(image_url or "").strip()
    if not url:
        return False

    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        return False

    req = Request(
        url,
        headers={
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36"
            )
        },
    )

    try:
        with urlopen(req, timeout=10) as response:
            chunks = []
            total = 0
            max_bytes = 20 * 1024 * 1024
            while True:
                chunk = response.read(64 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > max_bytes:
                    return False
                chunks.append(chunk)
            raw = b"".join(chunks)
    except OSError:
        return False

    if not raw:
        return False

    with open(out_path, "wb") as f:
        f.write(raw)
    return True


def _looks_like_supported_image_file(file_path):
    try:
        with open(file_path, "rb") as f:
            head = f.read(16)
    except OSError:
        return False

    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return True
    if head.startswith(b"\xFF\xD8\xFF"):
        return True
    if head.startswith(b"GIF87a") or head.startswith(b"GIF89a"):
        return True
    if head.startswith(b"RIFF") and b"WEBP" in head:
        return True
    return False


def _render_visual_figure(
    doc,
    item,
    asset_dir,
    index,
    is_submission=False,
    caption="",
):
    kind = item.get("kind") or "diagram"
    title = _clean_inline_text(item.get("title"), "시각자료")
    purpose = _clean_inline_text(item.get("purpose"))
    scoring = _clean_inline_text(item.get("scoringPoint"))

    label = caption or ("그래프" if kind == "graph" else "도해")
    if is_submission:
        add_content_sub(doc, f"{label}: {title}", indent=0.8)
    else:
        add_content_line(doc, 1, "", bold_prefix=f"[{label}] {title}: ")
    if purpose:
        add_content_sub(doc, f"목적: {purpose}", indent=1.0)

    image_path = asset_dir / f"visual_{index}_{kind}.png"
    image_data = item.get("imageData") or ""
    image_url = item.get("imageUrl") or ""
    decoded = _decode_image_data_to_file(image_data, str(image_path))
    downloaded = False
    if not decoded:
        downloaded = _download_image_from_url(image_url, str(image_path))

    if (
        (not decoded and not downloaded)
        or not _looks_like_supported_image_file(str(image_path))
    ):
        _generate_placeholder_visual_png(str(image_path), kind=kind)

    para = doc.add_paragraph()
    run = para.add_run()
    try:
        run.add_picture(str(image_path), width=Cm(15.5))
    except UnicodeDecodeError:
        _generate_placeholder_visual_png(str(image_path), kind=kind)
        run.add_picture(str(image_path), width=Cm(15.5))

    spec_lines = _clean_multiline_text(item.get("spec") or "", max_lines=4)
    for line in spec_lines:
        add_content_sub(doc, line, indent=1.0)
    if scoring:
        add_content_sub(doc, f"채점 포인트: {scoring}", indent=1.0)


def generate_bridge(payload):
    try:
        exam_no = payload.get("exam_no", 120)
        period = payload.get("period", 1)
        q_num = payload.get("q_num", 1)
        title = _clean_inline_text(
            payload.get("title", "Untitled Question"),
            "Untitled Question",
        )
        docx_style = str(
            payload.get("docx_style", "submission") or "submission"
        )
        is_submission = docx_style.lower() == "submission"
        raw_question = "\n".join(
            _clean_multiline_text(
                payload.get("raw_question", ""),
                max_lines=28,
            )
        ).strip()
        answer_text = "\n".join(
            _clean_multiline_text(
                payload.get("answer_text", ""),
                max_lines=32,
            )
        ).strip()
        llm_data = payload.get("llm_data", {})

        # 저장 경로 기준을 "현재 파일 기준 상대"로 고정하되,
        # 결과는 workspace 루트 하위 exported_subnotes 로 일관되게 배치한다.
        # (기존: solution/exported_subnotes)
        output_base_dir = WORKSPACE_ROOT / "exported_subnotes"
        output_base_dir.mkdir(parents=True, exist_ok=True)

        exam_dir = output_base_dir / f"{exam_no}회"
        exam_dir.mkdir(parents=True, exist_ok=True)
        asset_dir = exam_dir / "_assets"
        asset_dir.mkdir(parents=True, exist_ok=True)

        safe_title = re.sub(r'[\\/:*?"<>|\n\r]', '_', title).strip()
        safe_title = safe_title[:50]
        filename = f"{exam_no}_{period}_{q_num}_{safe_title}.docx"
        filepath = exam_dir / filename

        doc = create_document()
        add_exam_header(doc, q_num, title, exam_no=exam_no, period=period)

        if is_submission:
            answer_lines = _clean_multiline_text(answer_text, max_lines=28)
            if not answer_lines:
                fallback_lines = _clean_multiline_text(
                    llm_data.get("overview", "") or raw_question,
                    max_lines=8,
                )
                answer_lines = fallback_lines or [title]

            for line in answer_lines[:28]:
                add_content_sub(doc, line, indent=0.8)
        else:
            # 1. Overview
            add_section_header(doc, 1, "개요", "Overview")
            overview_text = llm_data.get("overview", title + "에 대한 정의입니다.")
            add_content_line(doc, 1, "", bold_prefix="정의: ")
            add_content_sub(doc, overview_text, indent=1.0)

            if raw_question:
                add_content_line(doc, 1, "", bold_prefix="문제 핵심: ")
                add_content_sub(doc, raw_question[:600], indent=1.0)

            if answer_text:
                add_content_line(doc, 1, "", bold_prefix="원문 모범답안 반영: ")
                answer_lines = _clean_multiline_text(answer_text, max_lines=18)
                if not answer_lines:
                    answer_lines = [answer_text]

                for line in answer_lines[:18]:
                    add_content_sub(doc, line, indent=1.0)

            # 2. Characteristics Table
            add_section_header(doc, 2, "주요 특징 및 분석", "Key Characteristics")
            col_w = [3.0, 5.0, TABLE_W - 3.0 - 5.0]
            chars = llm_data.get("characteristics", [])
            data_rows = []
            for c in chars:
                name = c.get("name", "특성")
                d1 = c.get("desc1", "설명")
                d2 = c.get("desc2", "-")
                data_rows.append([name, d1, d2])

            if not data_rows:
                data_rows = [["항목 A", "설명 A", "-"], ["항목 B", "설명 B", "-"]]

            add_data_table(
                doc,
                headers=["구 분", "특성 1", "특성 2"],
                data=data_rows,
                col_widths=col_w,
            )

            # 3. Professional Insights
            add_section_header(doc, 3, "기술사적 소견", "Professional Insight")
            insights = llm_data.get("insights", [])
            insight_rows = []
            for i in insights:
                insight_rows.append(
                    (i.get("title", "소견"), i.get("content", "-"))
                )

            if not insight_rows:
                insight_rows = [("실무 경험", "설계기준 준수"), ("최신 동향", "KDS 강화")]

            add_insight_box(doc, insight_rows)

        # 4. Visual Assets (diagram/table/graph/image)
        visual_items = _normalize_visual_items(llm_data, title)
        if visual_items:
            if not is_submission:
                add_section_header(
                    doc,
                    4,
                    "도해·표·그래프",
                    "Visual Assets for Submission",
                )
            counters = {"diagram": 0, "table": 0, "graph": 0}
            for idx, item in enumerate(visual_items, start=1):
                kind = (item.get("kind") or "diagram").lower()

                title_text = str(item.get("title") or "").lower()
                is_table_like_image = kind == "image" and bool(
                    re.search(r"비교표|table", title_text, re.IGNORECASE)
                )
                has_real_image = bool(
                    str(item.get("imageData") or "").strip()
                    or str(item.get("imageUrl") or "").strip()
                )

                if is_table_like_image and not has_real_image:
                    kind = "table"

                if kind == "table":
                    counters["table"] += 1
                    caption = f"표-{counters['table']}"
                    _render_visual_table(
                        doc,
                        item,
                        is_submission=is_submission,
                        caption=caption,
                    )
                elif kind in {"diagram", "graph", "image"}:
                    if kind == "graph":
                        counters["graph"] += 1
                        caption = f"그래프-{counters['graph']}"
                    elif is_table_like_image:
                        counters["table"] += 1
                        caption = f"표-{counters['table']}"
                    else:
                        counters["diagram"] += 1
                        caption = f"도해-{counters['diagram']}"
                    _render_visual_figure(
                        doc,
                        item,
                        asset_dir,
                        idx,
                        is_submission=is_submission,
                        caption=caption,
                    )
                else:
                    add_callout_box(
                        doc,
                        f"[VISUAL]: {item.get('title', '시각자료')}",
                        item.get("spec", ""),
                        bg_color="F0F9FF",
                        border_color="7DD3FC",
                    )

        if not is_submission:
            # Divider & Strategy
            add_divider(doc)

            keywords = llm_data.get("keywords", "키워드 미지정")
            strategy = llm_data.get("strategy", "문제 분석 기반 전략 수립")
            add_strategy_section(doc, [
                ("핵심 키워드", keywords),
                ("합격 포인트", strategy),
            ])

            add_footer(doc, q_num, title, exam_no=exam_no, period=period)

        save_document(doc, str(filepath))

        return {"ok": True, "path": str(filepath), "filename": filename}

    except (TypeError, ValueError, KeyError, OSError, RuntimeError):
        return {"ok": False, "error": traceback.format_exc()}


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(
            json.dumps(
                {"ok": False, "error": "No payload provided"},
                ensure_ascii=False,
            )
        )
        sys.exit(1)

    try:
        payload_data = json.loads(sys.argv[1])
        result = generate_bridge(payload_data)
        print(json.dumps(result, ensure_ascii=False))
    except (json.JSONDecodeError, TypeError, ValueError) as e:
        print(
            json.dumps(
                {"ok": False, "error": f"Main error: {e}"},
                ensure_ascii=False,
            )
        )
