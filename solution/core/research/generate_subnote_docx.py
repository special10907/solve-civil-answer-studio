#!/usr/bin/env python3
"""
Professional Structural Engineering PE Exam Sub-note Generator
Generates a high-quality .docx for Question 1: D-Region (응력 교란구역)
120th PE Exam, Session 1, Question 1
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


# ─── Configuration ───────────────────────────────────────────────────────────
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "120_1_1_Professional_Subnote.docx")

# Colors
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_NAVY = RGBColor(0x0D, 0x1B, 0x2A)
STEEL_BLUE = RGBColor(0x4A, 0x6F, 0xA5)
ACCENT_RED = RGBColor(0xC0, 0x39, 0x2B)
ACCENT_GOLD = RGBColor(0xD4, 0x8B, 0x0B)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY_BG = "F2F4F8"
HEADER_BG = "1B2A4A"
SUB_HEADER_BG = "E8ECF2"
SECTION_BG = "F7F9FC"
STRATEGY_BG = "FFF8E1"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with (size, color, style)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color, style) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = parse_xml(f'<w:{side} {nsdecls("w")} w:w="{val}" w:type="dxa"/>')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_run(paragraph, text, font_name="Malgun Gothic", size=10, bold=False,
            italic=False, color=DARK_GRAY, underline=False):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if underline:
        run.font.underline = True
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_horizontal_line(doc, color="1B2A4A", thickness=6):
    """Add a horizontal line using a 1-row table trick."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, color)
    # Set height
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{thickness}" w:hRule="exact"/>')
    trPr.append(trH)
    # Set width to full
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    # Remove existing tblW
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)
    return tbl


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(borders)


def set_table_width(table, width_pct=100):
    """Set table width as percentage."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_pct * 50}" w:type="pct"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)


def set_row_height(row, height_cm):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="exact"/>'
    )
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)
    trPr.append(trH)


def set_column_width(column, width_cm):
    """Set column width."""
    for cell in column.cells:
        cell.width = Cm(width_cm)


# ─── Document Creation ───────────────────────────────────────────────────────

doc = Document()

# ─── Page Setup (A4) ─────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Remove default paragraph spacing from Normal style
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION: HEADER BLOCK
# ═══════════════════════════════════════════════════════════════════════════════

header_table = doc.add_table(rows=3, cols=3)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(header_table, 100)

# --- Row 0: Main Title Bar ---
row0 = header_table.rows[0]
set_row_height(row0, 1.2)
for i in range(3):
    cell = row0.cells[i]
    set_cell_shading(cell, HEADER_BG)
    set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Merge row 0 cells
row0.cells[0].merge(row0.cells[2])
merged_cell = row0.cells[0]
merged_cell.text = ""
p = merged_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "[ 제 120 회 ]   ", "Arial", 13, bold=True, color=WHITE)
add_run(p, "토  목  구  조  기  술  사", "Malgun Gothic", 14, bold=True, color=WHITE)
add_run(p, "   ( 제 1 교 시 )", "Arial", 13, bold=True, color=WHITE)

# --- Row 1: Exam Info Bar ---
row1 = header_table.rows[1]
set_row_height(row1, 0.85)
for i in range(3):
    cell = row1.cells[i]
    set_cell_shading(cell, SUB_HEADER_BG)
    set_cell_margins(cell, top=40, bottom=40, left=120, right=120)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Cell 0: 성명
c = row1.cells[0]
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p, "성 명 :  ", "Malgun Gothic", 10, bold=True, color=NAVY)
add_run(p, "B  O  S  S", "Arial", 11, bold=True, color=ACCENT_RED)

# Cell 1: 문제번호
c = row1.cells[1]
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "문제번호 :  ", "Malgun Gothic", 10, bold=True, color=NAVY)
add_run(p, "1", "Arial", 12, bold=True, color=ACCENT_RED)

# Cell 2: 페이지
c = row1.cells[2]
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, "Page  ", "Arial", 9, color=MEDIUM_GRAY)
add_run(p, "1 / 1", "Arial", 9, bold=True, color=NAVY)

# --- Row 2: Question Title ---
row2 = header_table.rows[2]
set_row_height(row2, 1.0)
for i in range(3):
    cell = row2.cells[i]
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

row2.cells[0].merge(row2.cells[2])
merged_cell = row2.cells[0]
merged_cell.text = ""
p = merged_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p, "문 1.  ", "Malgun Gothic", 12, bold=True, color=DARK_NAVY)
add_run(p, "철근콘크리트 보의 응력 교란구역 (D-Region)", "Malgun Gothic", 12, bold=True, color=DARK_NAVY)

# Set header table borders
tbl = header_table._tbl
tblPr = tbl.tblPr
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:right w:val="single" w:sz="12" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D5E0"/>'
    f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
for existing in tblPr.findall(qn('w:tblBorders')):
    tblPr.remove(existing)
tblPr.append(borders)

# Spacer
doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPER: Add Section Header
# ═══════════════════════════════════════════════════════════════════════════════

def add_section_header(doc, number, title_kr, title_en, accent_color=NAVY):
    """Add a professional section header with number badge."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(tbl, 100)
    remove_table_borders(tbl)

    # Number badge cell
    badge_cell = tbl.cell(0, 0)
    badge_cell.width = Cm(0.9)
    set_cell_shading(badge_cell, accent_color.hex_string if hasattr(accent_color, 'hex_string') else HEADER_BG)
    badge_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(badge_cell, top=40, bottom=40, left=60, right=60)
    p = badge_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(number), "Arial", 11, bold=True, color=WHITE)

    # Title cell
    title_cell = tbl.cell(0, 1)
    set_cell_margins(title_cell, top=40, bottom=40, left=120, right=60)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Add bottom border to title cell
    set_cell_border(title_cell, bottom=(8, accent_color.hex_string if hasattr(accent_color, 'hex_string') else HEADER_BG, "single"))
    p = title_cell.paragraphs[0]
    add_run(p, f"{title_kr} ", "Malgun Gothic", 11, bold=True, color=DARK_NAVY)
    add_run(p, f"({title_en})", "Arial", 9, italic=True, color=STEEL_BLUE)

    # Set specific column widths
    # Badge column
    for cell in tbl.columns[0].cells:
        cell.width = Cm(0.9)
    # Ensure no cell margins override
    return tbl


def add_content_line(doc, number, text, bold_prefix="", is_sub=False, indent_level=0):
    """Add a numbered content line."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line_spacing=1.3)
    left_indent = 0.5 + (indent_level * 0.5)
    p.paragraph_format.left_indent = Cm(left_indent)

    markers = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩"]
    if is_sub:
        markers = ["▸", "▸", "▸", "▸", "▸"]

    marker = markers[min(number - 1, len(markers) - 1)] if number > 0 else "•"

    add_run(p, f" {marker} ", "Arial", 9, color=STEEL_BLUE)
    if bold_prefix:
        add_run(p, bold_prefix, "Malgun Gothic", 10, bold=True, color=DARK_NAVY)
    add_run(p, text, "Malgun Gothic", 10, color=DARK_GRAY)
    return p


def add_keyword_badge(paragraph, text, bg_color=SUB_HEADER_BG, text_color=NAVY):
    """Add an inline keyword badge (simulated with bold colored text)."""
    add_run(paragraph, f" [{text}] ", "Arial", 9, bold=True, color=text_color)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: 개요
# ═══════════════════════════════════════════════════════════════════════════════

add_section_header(doc, 1, "개요", "Definition & Bernoulli Hypothesis")

add_content_line(doc, 1, "", bold_prefix="개념: ")
p = doc.add_paragraph()
set_paragraph_spacing(p, before=0, after=2, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(1.2)
add_run(p, "베르누이의 평면유지 가설(Bernoulli Hypothesis)이 성립하지 않는 ", "Malgun Gothic", 10, color=DARK_GRAY)
add_run(p, "비선형 변형률 분포 구간", "Malgun Gothic", 10, bold=True, color=ACCENT_RED)
add_run(p, ".", "Malgun Gothic", 10, color=DARK_GRAY)

add_content_line(doc, 2, "", bold_prefix="발생: ")
p = doc.add_paragraph()
set_paragraph_spacing(p, before=0, after=2, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(1.2)
add_run(p, "하중의 집중(Statical) 또는 형상의 급변(Geometric)으로 응력 흐름이 교란됨.", "Malgun Gothic", 10, color=DARK_GRAY)

add_content_line(doc, 3, "", bold_prefix="설계: ")
p = doc.add_paragraph()
set_paragraph_spacing(p, before=0, after=4, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(1.2)
add_run(p, "일반 휨 이론 적용 불가 → ", "Malgun Gothic", 10, color=DARK_GRAY)
add_run(p, "STM(Strut-and-Tie Model) 설계 필수", "Malgun Gothic", 10, bold=True, color=NAVY)
add_run(p, ".", "Malgun Gothic", 10, color=DARK_GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: D-Region 발생 원인 및 범위
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_paragraph_spacing(p, before=6, after=0)

add_section_header(doc, 2, "D-Region의 발생 원인 및 범위", "Saint-Venant's Principle")

# Cause table
cause_table = doc.add_table(rows=4, cols=3)
cause_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(cause_table, 95)

# Header row
headers = ["구 분", "원 인", "사 례"]
for i, h in enumerate(headers):
    cell = cause_table.cell(0, i)
    set_cell_shading(cell, HEADER_BG)
    set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, "Malgun Gothic", 9, bold=True, color=WHITE)

# Data rows
data = [
    ["기하학적 불연속\n(Geometric)", "형상의 급변", "개구부, 댑트 보(Dapped-end),\n단면 급변부"],
    ["하중 불연속\n(Statical)", "하중의 집중", "집중하중 작용점,\n지점 반력 작용부"],
    ["범위 산정\n(Range)", "Saint-Venant 원리", "불연속점으로부터\n부재 깊이(h) 이내"],
]

for row_idx, row_data in enumerate(data):
    bg = LIGHT_GRAY_BG if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, text in enumerate(row_data):
        cell = cause_table.cell(row_idx + 1, col_idx)
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        if col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, "Malgun Gothic", 9, bold=True, color=NAVY)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, "Malgun Gothic", 9, color=DARK_GRAY)

# Table borders
tbl_el = cause_table._tbl
tblPr_el = tbl_el.tblPr
borders_el = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D5E0"/>'
    f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D5E0"/>'
    f'</w:tblBorders>'
)
for existing in tblPr_el.findall(qn('w:tblBorders')):
    tblPr_el.remove(existing)
tblPr_el.append(borders_el)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: 실전 삽도 (Professional Diagram using Tables)
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_paragraph_spacing(p, before=6, after=0)

add_section_header(doc, 3, "실전 삽도", "Hand-drawable Sketch for Exam")

# Instruction box
inst_tbl = doc.add_table(rows=1, cols=1)
inst_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(inst_tbl, 95)
cell = inst_tbl.cell(0, 0)
set_cell_shading(cell, "FFF3E0")
set_cell_margins(cell, top=60, bottom=60, left=120, right=120)

p = cell.paragraphs[0]
add_run(p, "📝 시험장 드로잉 전략", "Malgun Gothic", 9, bold=True, color=RGBColor(0xE6, 0x5C, 0x00))
p2 = cell.add_paragraph()
set_paragraph_spacing(p2, before=2, after=0, line_spacing=1.3)
steps = [
    "단순보 외곽선을 그린다.",
    "집중하중 P와 지점부 R을 표시한다.",
    "각 작용점으로부터 수평거리 h만큼을 점선으로 구획한다.",
    "구획된 영역에 'D-Region', 그 외 영역에 'B-Region'을 명기한다.",
    "D-Region 내부에 Strut(압축대)를 대각선으로 그려 넣는다."
]
for i, step in enumerate(steps):
    add_run(p2, f"  {i+1}. {step}", "Malgun Gothic", 8.5, color=DARK_GRAY)
    if i < len(steps) - 1:
        add_run(p2, "\n", "Malgun Gothic", 8.5)

# Border for instruction box
tbl_el2 = inst_tbl._tbl
tblPr_el2 = tbl_el2.tblPr
brd = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'</w:tblBorders>'
)
for existing in tblPr_el2.findall(qn('w:tblBorders')):
    tblPr_el2.remove(existing)
tblPr_el2.append(brd)

# Spacer
p = doc.add_paragraph()
set_paragraph_spacing(p, before=4, after=0)

# ─── Professional Diagram (Table-based) ──────────────────────────────────────
# Main beam diagram using nested tables

diagram_table = doc.add_table(rows=5, cols=5)
diagram_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(diagram_table, 85)
remove_table_borders(diagram_table)

# Row 0: Load indicator - "P (집중하중)" centered above D1
set_row_height(diagram_table.rows[0], 0.7)
for i in range(5):
    c = diagram_table.cell(0, i)
    c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    set_cell_margins(c, top=0, bottom=0, left=20, right=20)

# Merge center cells for load arrow
diagram_table.cell(0, 1).merge(diagram_table.cell(0, 3))
center_load = diagram_table.cell(0, 1)
center_load.text = ""
p = center_load.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "P (집중하중)", "Arial", 9, bold=True, color=ACCENT_RED)

# Row 1: Arrow
set_row_height(diagram_table.rows[1], 0.4)
diagram_table.cell(1, 1).merge(diagram_table.cell(1, 3))
arrow_cell = diagram_table.cell(1, 1)
arrow_cell.text = ""
p = arrow_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "▼", "Arial", 14, bold=True, color=ACCENT_RED)

# Row 2: Beam body - D1 | B-Region | D2 with dimensions
set_row_height(diagram_table.rows[2], 1.5)

# D1
d1_cell = diagram_table.cell(2, 0)
d1_cell.merge(diagram_table.cell(2, 1))
d1 = diagram_table.cell(2, 0)
set_cell_shading(d1, "DCEEFB")
set_cell_margins(d1, top=30, bottom=30, left=40, right=40)
d1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_border(d1,
    top=(12, "1B2A4A", "single"),
    bottom=(12, "1B2A4A", "single"),
    left=(12, "1B2A4A", "single"),
    right=(6, "4A6FA5", "dashed"))
p = d1.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "D-Region\n(D1)", "Arial", 9, bold=True, color=NAVY)
p2 = d1.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, "비선형 응력", "Malgun Gothic", 7.5, italic=True, color=STEEL_BLUE)

# B
b_cell = diagram_table.cell(2, 2)
set_cell_shading(b_cell, "F7F9FC")
set_cell_margins(b_cell, top=30, bottom=30, left=40, right=40)
b_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_border(b_cell,
    top=(12, "1B2A4A", "single"),
    bottom=(12, "1B2A4A", "single"),
    left=(6, "4A6FA5", "dashed"),
    right=(6, "4A6FA5", "dashed"))
p = b_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "B-Region", "Arial", 9, bold=True, color=DARK_GRAY)
p2 = b_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, "선형 응력\n(평면유지)", "Malgun Gothic", 7.5, italic=True, color=MEDIUM_GRAY)

# D2
d2_cell = diagram_table.cell(2, 3)
d2_cell.merge(diagram_table.cell(2, 4))
d2 = diagram_table.cell(2, 3)
set_cell_shading(d2, "DCEEFB")
set_cell_margins(d2, top=30, bottom=30, left=40, right=40)
d2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_border(d2,
    top=(12, "1B2A4A", "single"),
    bottom=(12, "1B2A4A", "single"),
    left=(6, "4A6FA5", "dashed"),
    right=(12, "1B2A4A", "single"))
p = d2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "D-Region\n(D2)", "Arial", 9, bold=True, color=NAVY)
p2 = d2.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, "비선형 응력", "Malgun Gothic", 7.5, italic=True, color=STEEL_BLUE)

# Row 3: Support indicators
set_row_height(diagram_table.rows[3], 0.5)
for i in range(5):
    c = diagram_table.cell(3, i)
    set_cell_margins(c, top=0, bottom=0, left=20, right=20)
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# Left support
sup_left = diagram_table.cell(3, 0)
sup_left.text = ""
p = sup_left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "▲", "Arial", 14, bold=True, color=DARK_NAVY)

# Right support
sup_right = diagram_table.cell(3, 4)
sup_right.text = ""
p = sup_right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "▲", "Arial", 14, bold=True, color=DARK_NAVY)

# Row 4: Labels
set_row_height(diagram_table.rows[4], 0.6)
for i in range(5):
    c = diagram_table.cell(4, i)
    set_cell_margins(c, top=0, bottom=0, left=20, right=20)
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# Left label
l_label = diagram_table.cell(4, 0)
l_label.text = ""
p = l_label.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "지점 R₁", "Malgun Gothic", 8, color=MEDIUM_GRAY)

# Center dimension label
diagram_table.cell(4, 1).merge(diagram_table.cell(4, 3))
dim_label = diagram_table.cell(4, 1)
dim_label.text = ""
p = dim_label.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "│← h →│← B-Region →│← h →│", "Arial", 8, color=STEEL_BLUE)

# Right label
r_label = diagram_table.cell(4, 4)
r_label.text = ""
p = r_label.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "지점 R₂", "Malgun Gothic", 8, color=MEDIUM_GRAY)

# Dimension annotation on the right side (h = depth)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_paragraph_spacing(p, before=2, after=2, line_spacing=1.0)
add_run(p, "※ h = 부재 깊이(Depth)  |  D = Disturbed  |  B = Bernoulli(건전구역)", "Malgun Gothic", 8, italic=True, color=MEDIUM_GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: STM 설계 메커니즘
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_paragraph_spacing(p, before=6, after=0)

add_section_header(doc, 4, "STM 설계 메커니즘", "Strut-and-Tie Model / KDS 14 20:2024")

# STM Components Table
stm_table = doc.add_table(rows=4, cols=4)
stm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(stm_table, 95)

# Header
stm_headers = ["구성요소", "역 할", "설계 핵심", "Node 유형"]
for i, h in enumerate(stm_headers):
    cell = stm_table.cell(0, i)
    set_cell_shading(cell, HEADER_BG)
    set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, "Malgun Gothic", 9, bold=True, color=WHITE)

# STM data
stm_data = [
    ["Strut\n(압축대)", "콘크리트 압축 응력장\n하중→절점 전달", "유효 압축 강도 검토\n(Bottle/Prismatic형)", "—"],
    ["Tie\n(인장재)", "철근의 인장 응력 수용\n및 정착", "항복 강도 확보\n정착 길이 검토", "—"],
    ["Node\n(절점)", "Strut+Tie 교차부\n하중 평형 구역", "지압 강도 검토\n정착 상세 확보", "CCC, CCT,\nCTT"],
]

for row_idx, row_data in enumerate(stm_data):
    bg = LIGHT_GRAY_BG if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, text in enumerate(row_data):
        cell = stm_table.cell(row_idx + 1, col_idx)
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        if col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, "Malgun Gothic", 9, bold=True, color=NAVY)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, "Malgun Gothic", 8.5, color=DARK_GRAY)

# Table borders
tbl_el3 = stm_table._tbl
tblPr_el3 = tbl_el3.tblPr
brd3 = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D5E0"/>'
    f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D5E0"/>'
    f'</w:tblBorders>'
)
for existing in tblPr_el3.findall(qn('w:tblBorders')):
    tblPr_el3.remove(existing)
tblPr_el3.append(brd3)

# Design core note
p = doc.add_paragraph()
set_paragraph_spacing(p, before=4, after=2, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(0.5)
add_run(p, "▶ 설계 핵심: ", "Malgun Gothic", 9, bold=True, color=NAVY)
add_run(p, "하중 경로(Load Path) 시각화 및 절점부 지압/정착 상세 검토 (KDS 14 20:2024 준수)", "Malgun Gothic", 9, color=DARK_GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: 기술사적 소견
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_paragraph_spacing(p, before=6, after=0)

add_section_header(doc, 5, "기술사적 소견", "Professional Insight")

# Insight box
insight_tbl = doc.add_table(rows=1, cols=1)
insight_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(insight_tbl, 95)
cell = insight_tbl.cell(0, 0)
set_cell_shading(cell, "F0F4FF")
set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

insights = [
    ("실무 경험 (17년):", "D-Region 설계 실패는 대개 '철근의 정착' 및 '절점부의 압축 압쇄'에서 기인함."),
    ("설계 착안점:", "전단 설계 시 단순 Vc+Vs 식에 의존하지 말고, STM을 통한 하중의 물리적 거동 파악이 기술사의 핵심 역량임."),
    ("최신 경향:", "2024 KDS 기준은 복잡한 불연속 구간에서 STM 설계를 원칙으로 하며, 특히 절점 상세 설계(Detailing)를 강조함."),
]

for i, (title, desc) in enumerate(insights):
    p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
    set_paragraph_spacing(p, before=2, after=4, line_spacing=1.35)
    add_run(p, f"  ■ {title} ", "Malgun Gothic", 9, bold=True, color=NAVY)
    add_run(p, desc, "Malgun Gothic", 9, color=DARK_GRAY)

# Border for insight box
tbl_el4 = insight_tbl._tbl
tblPr_el4 = tbl_el4.tblPr
brd4 = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="4A6FA5"/>'
    f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="1B2A4A"/>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="4A6FA5"/>'
    f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="4A6FA5"/>'
    f'</w:tblBorders>'
)
for existing in tblPr_el4.findall(qn('w:tblBorders')):
    tblPr_el4.remove(existing)
tblPr_el4.append(brd4)


# ═══════════════════════════════════════════════════════════════════════════════
# DIVIDER: 이하 빈칸
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_paragraph_spacing(p, before=8, after=0)

add_horizontal_line(doc, "D0D5E0", 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=4, after=4)
add_run(p, "──────────────  이  하  빈  칸  ──────────────", "Malgun Gothic", 10, color=MEDIUM_GRAY)

add_horizontal_line(doc, "D0D5E0", 12)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION: 문제 분석 및 핵심 전략 (Strategy Box)
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_paragraph_spacing(p, before=8, after=0)

# Strategy header
strat_header = doc.add_table(rows=1, cols=1)
strat_header.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(strat_header, 100)
cell = strat_header.cell(0, 0)
set_cell_shading(cell, "D48B0B")
set_cell_margins(cell, top=50, bottom=50, left=120, right=120)
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "■  문제 분석 및 핵심 전략  ■", "Malgun Gothic", 11, bold=True, color=WHITE)

# Remove borders from header
remove_table_borders(strat_header)

# Strategy content table
strat_tbl = doc.add_table(rows=4, cols=2)
strat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(strat_tbl, 100)

strategy_items = [
    ("핵심 키워드", "Saint-Venant's Principle  |  STM (Strut-and-Tie Model)  |  Bernoulli Hypothesis  |  h (Depth)"),
    ("반드시 포함", "① 평면유지 가설 파괴 개념  ② D-Region 범위 설정 기준 (부재깊이 h)  ③ STM 구성요소 및 설계 핵심"),
    ("유사 기출문제", "• 108회 1-5 (개구부 설계)  • 125회 2-3 (댑트 보 계산)  • 132회 1-8 (지압 설계)"),
    ("합격 포인트", "단순 암기 답안이 아닌, '하중 경로(Load Path)'의 이해를 보여주는 삽도 배치가 당락을 결정함."),
]

for idx, (label, content) in enumerate(strategy_items):
    bg = STRATEGY_BG if idx % 2 == 0 else "FFFFFF"

    # Label cell
    label_cell = strat_tbl.cell(idx, 0)
    set_cell_shading(label_cell, "FFF3E0")
    set_cell_margins(label_cell, top=50, bottom=50, left=80, right=80)
    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    label_cell.width = Cm(3.0)
    p = label_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, label, "Malgun Gothic", 9, bold=True, color=RGBColor(0xD4, 0x8B, 0x0B))

    # Content cell
    content_cell = strat_tbl.cell(idx, 1)
    set_cell_shading(content_cell, bg)
    set_cell_margins(content_cell, top=50, bottom=50, left=100, right=80)
    content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = content_cell.paragraphs[0]
    add_run(p, content, "Malgun Gothic", 9, color=DARK_GRAY)

# Strategy table borders
tbl_el5 = strat_tbl._tbl
tblPr_el5 = tbl_el5.tblPr
brd5 = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="D48B0B"/>'
    f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="D48B0B"/>'
    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="D48B0B"/>'
    f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="D48B0B"/>'
    f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="F5DEB3"/>'
    f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="F5DEB3"/>'
    f'</w:tblBorders>'
)
for existing in tblPr_el5.findall(qn('w:tblBorders')):
    tblPr_el5.remove(existing)
tblPr_el5.append(brd5)


# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_paragraph_spacing(p, before=12, after=0)

add_horizontal_line(doc, HEADER_BG, 8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=4, after=0)
add_run(p, "BOSS  |  제120회 토목구조기술사  |  1교시 1번  |  D-Region & STM", "Malgun Gothic", 8, color=MEDIUM_GRAY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p2, before=2, after=0)
add_run(p2, "Generated by flowithOS  ·  Professional PE Exam Sub-note System", "Arial", 7, italic=True, color=RGBColor(0xAA, 0xAA, 0xAA))


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_FILE)
print(f"✅ Document generated successfully: {OUTPUT_FILE}")
print(f"   File size: {os.path.getsize(OUTPUT_FILE):,} bytes")
