#!/usr/bin/env python3
"""
Professional Structural Engineering PE Exam Sub-note Generator v2
Refined layout: optimized column widths, tighter spacing, visual polish.
120th PE Exam, Session 1, Question 1 — D-Region (응력 교란구역)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Configuration ───────────────────────────────────────────────────────────
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "120_1_1_Professional_Subnote_v2.docx")

# Available content width: A4(21cm) - 2cm left - 2cm right = 17cm
CONTENT_WIDTH_CM = 17.0

# Colors
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_NAVY = RGBColor(0x0D, 0x1B, 0x2A)
STEEL_BLUE = RGBColor(0x4A, 0x6F, 0xA5)
ACCENT_RED = RGBColor(0xC0, 0x39, 0x2B)
ACCENT_GOLD = RGBColor(0xD4, 0x8B, 0x0B)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Background hex strings
LIGHT_GRAY_BG = "F2F4F8"
HEADER_BG = "1B2A4A"
SUB_HEADER_BG = "E8ECF2"
SECTION_BG = "F7F9FC"
STRATEGY_BG = "FFF8E1"
GOLD_HEX = "D48B0B"


# ─── Utility functions ───────────────────────────────────────────────────────

def cm_to_twips(cm):
    """Convert cm to twips (1 cm = 567 twips)."""
    return int(cm * 567)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color, style) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = parse_xml(f'<w:{side} {nsdecls("w")} w:w="{val}" w:type="dxa"/>')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_cell_width(cell, width_cm):
    """Set explicit cell width via tcW."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tw = cm_to_twips(width_cm)
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{tw}" w:type="dxa"/>')
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcPr.append(tcW)


def set_table_col_widths(table, widths_cm):
    """Set explicit column widths via tblGrid and individual cell widths."""
    tbl = table._tbl
    # Build tblGrid
    tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
    for w in widths_cm:
        gridCol = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{cm_to_twips(w)}"/>')
        tblGrid.append(gridCol)
    # Remove existing tblGrid
    for existing in tbl.findall(qn('w:tblGrid')):
        tbl.remove(existing)
    # Insert after tblPr
    tbl.insert(tbl.index(tbl.tblPr) + 1, tblGrid)
    # Also set each cell width for consistency
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                set_cell_width(row.cells[i], w)


def add_run(paragraph, text, font_name="Malgun Gothic", size=10, bold=False,
            italic=False, color=DARK_GRAY, underline=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if underline:
        run.font.underline = True
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_horizontal_line(doc, color="1B2A4A", thickness=6):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, color)
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{thickness}" w:hRule="exact"/>')
    trPr.append(trH)
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)
    return tbl


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(borders)


def set_table_width_pct(table, width_pct=100):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_pct * 50}" w:type="pct"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)


def set_table_width_fixed(table, width_cm):
    """Set table width as fixed twip value."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tw = cm_to_twips(width_cm)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{tw}" w:type="dxa"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)
    # Disable autofit so fixed widths are respected
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblPr.append(tblLayout)


def set_row_height(row, height_cm, rule="exact"):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{cm_to_twips(height_cm)}" w:hRule="{rule}"/>'
    )
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)
    trPr.append(trH)


def apply_table_borders(table, outer_color, inner_h_color="D0D5E0", inner_v_color="D0D5E0",
                         outer_sz=8, inner_sz=4):
    tbl_el = table._tbl
    tblPr_el = tbl_el.tblPr
    brd = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:left w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:bottom w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:right w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:insideH w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_h_color}"/>'
        f'  <w:insideV w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_v_color}"/>'
        f'</w:tblBorders>'
    )
    for existing in tblPr_el.findall(qn('w:tblBorders')):
        tblPr_el.remove(existing)
    tblPr_el.append(brd)


def add_spacer(doc, pts=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=pts, after=0)
    return p


# ─── Section Header Helper ──────────────────────────────────────────────────

BADGE_WIDTH = 0.7  # cm — compact badge
TITLE_WIDTH = CONTENT_WIDTH_CM - BADGE_WIDTH  # rest of the line

def add_section_header(doc, number, title_kr, title_en):
    """Compact numbered section header with underline."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width_fixed(tbl, CONTENT_WIDTH_CM)
    set_table_col_widths(tbl, [BADGE_WIDTH, TITLE_WIDTH])
    remove_table_borders(tbl)

    # Badge
    badge = tbl.cell(0, 0)
    set_cell_shading(badge, HEADER_BG)
    badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(badge, top=35, bottom=35, left=0, right=0)
    p = badge.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(number), "Arial", 11, bold=True, color=WHITE)

    # Title
    title = tbl.cell(0, 1)
    set_cell_margins(title, top=35, bottom=35, left=100, right=40)
    title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(title, bottom=(8, HEADER_BG, "single"))
    p = title.paragraphs[0]
    add_run(p, f"{title_kr} ", "Malgun Gothic", 11, bold=True, color=DARK_NAVY)
    add_run(p, f"({title_en})", "Arial", 9, italic=True, color=STEEL_BLUE)
    return tbl


def add_content_line(doc, number, text, bold_prefix="", indent_level=0):
    """Numbered content line with circled number marker."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line_spacing=1.3)
    p.paragraph_format.left_indent = Cm(0.4 + (indent_level * 0.4))
    markers = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩"]
    marker = markers[min(number - 1, len(markers) - 1)] if number > 0 else "•"
    add_run(p, f" {marker} ", "Arial", 9, color=STEEL_BLUE)
    if bold_prefix:
        add_run(p, bold_prefix, "Malgun Gothic", 10, bold=True, color=DARK_NAVY)
    add_run(p, text, "Malgun Gothic", 10, color=DARK_GRAY)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CREATION
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page Setup (A4)
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ═══════════════════════════════════════════════════════════════════════════════
# HEADER BLOCK — 3 rows with explicit column sizing
# ═══════════════════════════════════════════════════════════════════════════════

header_table = doc.add_table(rows=3, cols=3)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(header_table, CONTENT_WIDTH_CM)
# Row 1 (info bar) columns: name(7cm) | question(5cm) | page(5cm)
set_table_col_widths(header_table, [7.0, 5.0, 5.0])

# --- Row 0: Title Bar (merged) ---
row0 = header_table.rows[0]
set_row_height(row0, 1.1)
for i in range(3):
    cell = row0.cells[i]
    set_cell_shading(cell, HEADER_BG)
    set_cell_margins(cell, top=50, bottom=50, left=100, right=100)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

row0.cells[0].merge(row0.cells[2])
merged = row0.cells[0]
merged.text = ""
p = merged.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "[ 제 120 회 ]   ", "Arial", 13, bold=True, color=WHITE)
add_run(p, "토  목  구  조  기  술  사", "Malgun Gothic", 14, bold=True, color=WHITE)
add_run(p, "   ( 제 1 교 시 )", "Arial", 13, bold=True, color=WHITE)

# --- Row 1: Info Bar ---
row1 = header_table.rows[1]
set_row_height(row1, 0.8)
for i in range(3):
    cell = row1.cells[i]
    set_cell_shading(cell, SUB_HEADER_BG)
    set_cell_margins(cell, top=35, bottom=35, left=100, right=100)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Name
c = row1.cells[0]
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p, "성 명 :  ", "Malgun Gothic", 10, bold=True, color=NAVY)
add_run(p, "B  O  S  S", "Arial", 11, bold=True, color=ACCENT_RED)

# Question number
c = row1.cells[1]
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "문제번호 :  ", "Malgun Gothic", 10, bold=True, color=NAVY)
add_run(p, "1", "Arial", 12, bold=True, color=ACCENT_RED)

# Page
c = row1.cells[2]
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, "Page  ", "Arial", 9, color=MEDIUM_GRAY)
add_run(p, "1 / 1", "Arial", 9, bold=True, color=NAVY)

# --- Row 2: Question Title (merged) ---
row2 = header_table.rows[2]
set_row_height(row2, 0.9)
for i in range(3):
    cell = row2.cells[i]
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

row2.cells[0].merge(row2.cells[2])
merged = row2.cells[0]
merged.text = ""
p = merged.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p, "문 1.  ", "Malgun Gothic", 12, bold=True, color=DARK_NAVY)
add_run(p, "철근콘크리트 보의 응력 교란구역 (D-Region)", "Malgun Gothic", 12, bold=True, color=DARK_NAVY)

# Header borders
apply_table_borders(header_table, HEADER_BG, inner_h_color="D0D5E0", inner_v_color="D0D5E0",
                     outer_sz=12, inner_sz=4)
# Override insideV to none for merged rows
tbl_pr = header_table._tbl.tblPr
for brd in tbl_pr.findall(qn('w:tblBorders')):
    for iv in brd.findall(qn('w:insideV')):
        iv.set(qn('w:val'), 'none')

add_spacer(doc, 5)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: 개요
# ═══════════════════════════════════════════════════════════════════════════════

add_section_header(doc, 1, "개요", "Definition & Bernoulli Hypothesis")

add_content_line(doc, 1, "", bold_prefix="개념: ")
p = doc.add_paragraph()
set_paragraph_spacing(p, before=0, after=2, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(1.0)
add_run(p, "베르누이의 평면유지 가설(Bernoulli Hypothesis)이 성립하지 않는 ", "Malgun Gothic", 10, color=DARK_GRAY)
add_run(p, "비선형 변형률 분포 구간", "Malgun Gothic", 10, bold=True, color=ACCENT_RED)
add_run(p, ".", "Malgun Gothic", 10, color=DARK_GRAY)

add_content_line(doc, 2, "", bold_prefix="발생: ")
p = doc.add_paragraph()
set_paragraph_spacing(p, before=0, after=2, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(1.0)
add_run(p, "하중의 집중(Statical) 또는 형상의 급변(Geometric)으로 응력 흐름이 교란됨.", "Malgun Gothic", 10, color=DARK_GRAY)

add_content_line(doc, 3, "", bold_prefix="설계: ")
p = doc.add_paragraph()
set_paragraph_spacing(p, before=0, after=3, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(1.0)
add_run(p, "일반 휨 이론 적용 불가 → ", "Malgun Gothic", 10, color=DARK_GRAY)
add_run(p, "STM(Strut-and-Tie Model) 설계 필수", "Malgun Gothic", 10, bold=True, color=NAVY)
add_run(p, ".", "Malgun Gothic", 10, color=DARK_GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: D-Region 발생 원인 및 범위
# Column layout: Label(3.0cm) | Cause(4.5cm) | Examples(8.6cm) = ~16.1cm (95%)
# ═══════════════════════════════════════════════════════════════════════════════

add_spacer(doc, 5)
add_section_header(doc, 2, "D-Region의 발생 원인 및 범위", "Saint-Venant's Principle")

TABLE2_W = CONTENT_WIDTH_CM * 0.95  # ~16.15cm
COL2_WIDTHS = [3.0, 4.5, TABLE2_W - 3.0 - 4.5]  # [3.0, 4.5, 8.65]

cause_table = doc.add_table(rows=4, cols=3)
cause_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(cause_table, TABLE2_W)
set_table_col_widths(cause_table, COL2_WIDTHS)

# Header row
headers2 = ["구  분", "원  인", "사  례"]
for i, h in enumerate(headers2):
    cell = cause_table.cell(0, i)
    set_cell_shading(cell, HEADER_BG)
    set_cell_margins(cell, top=45, bottom=45, left=60, right=60)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, "Malgun Gothic", 9, bold=True, color=WHITE)

# Data
cause_data = [
    ["기하학적 불연속\n(Geometric)", "형상의 급변", "개구부, 댑트 보(Dapped-end), 단면 급변부"],
    ["하중 불연속\n(Statical)", "하중의 집중", "집중하중 작용점, 지점 반력 작용부"],
    ["범위 산정\n(Range)", "Saint-Venant 원리", "불연속점으로부터 부재 깊이(h) 이내"],
]

for row_idx, row_data in enumerate(cause_data):
    bg = LIGHT_GRAY_BG if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, text in enumerate(row_data):
        cell = cause_table.cell(row_idx + 1, col_idx)
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        if col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, "Malgun Gothic", 9, bold=True, color=NAVY)
        elif col_idx == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, "Malgun Gothic", 9, color=DARK_GRAY)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, "Malgun Gothic", 9, color=DARK_GRAY)

apply_table_borders(cause_table, HEADER_BG)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: 실전 삽도
# ═══════════════════════════════════════════════════════════════════════════════

add_spacer(doc, 5)
add_section_header(doc, 3, "실전 삽도", "Hand-drawable Sketch for Exam")

# Drawing instruction box — full width
inst_tbl = doc.add_table(rows=1, cols=1)
inst_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(inst_tbl, TABLE2_W)
cell = inst_tbl.cell(0, 0)
set_cell_shading(cell, "FFF3E0")
set_cell_margins(cell, top=50, bottom=50, left=100, right=100)

p = cell.paragraphs[0]
set_paragraph_spacing(p, before=0, after=0, line_spacing=1.2)
add_run(p, "  시험장 드로잉 전략  ", "Malgun Gothic", 9, bold=True, color=RGBColor(0xE6, 0x5C, 0x00))
p2 = cell.add_paragraph()
set_paragraph_spacing(p2, before=2, after=0, line_spacing=1.25)
steps = [
    "단순보 외곽선 → ",
    "집중하중 P, 지점 R 표시 → ",
    "작용점에서 h만큼 점선 구획 → ",
    "'D/B-Region' 명기 → ",
    "D-Region 내 Strut 대각선 삽입"
]
add_run(p2, "  " + " ".join(f"{i+1}){s}" for i, s in enumerate(steps)),
        "Malgun Gothic", 8, color=DARK_GRAY)

# Instruction box border
tbl_el2 = inst_tbl._tbl
tblPr_el2 = tbl_el2.tblPr
brd = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="FFB74D"/>'
    f'</w:tblBorders>'
)
for existing in tblPr_el2.findall(qn('w:tblBorders')):
    tblPr_el2.remove(existing)
tblPr_el2.append(brd)

add_spacer(doc, 3)

# ─── Beam Diagram (table-based, 5x5 with explicit widths) ────────────────────
# Total width ~14.5cm (85% of 17cm). Columns: D1a | D1b | B | D2a | D2b
DIAG_W = CONTENT_WIDTH_CM * 0.85  # ~14.45
DIAG_COL = DIAG_W / 5  # ~2.89 each

diagram_table = doc.add_table(rows=5, cols=5)
diagram_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(diagram_table, DIAG_W)
set_table_col_widths(diagram_table, [DIAG_COL] * 5)
remove_table_borders(diagram_table)

# Row 0: Load label
set_row_height(diagram_table.rows[0], 0.6, "atLeast")
for i in range(5):
    c = diagram_table.cell(0, i)
    c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    set_cell_margins(c, top=0, bottom=0, left=10, right=10)

diagram_table.cell(0, 1).merge(diagram_table.cell(0, 3))
cl = diagram_table.cell(0, 1)
cl.text = ""
p = cl.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "P (집중하중)", "Arial", 9, bold=True, color=ACCENT_RED)

# Row 1: Arrow
set_row_height(diagram_table.rows[1], 0.4, "atLeast")
diagram_table.cell(1, 1).merge(diagram_table.cell(1, 3))
ac = diagram_table.cell(1, 1)
ac.text = ""
p = ac.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "▼", "Arial", 14, bold=True, color=ACCENT_RED)

# Row 2: Beam body — D1(col0+1) | B(col2) | D2(col3+4)
set_row_height(diagram_table.rows[2], 1.4, "atLeast")

# D1
diagram_table.cell(2, 0).merge(diagram_table.cell(2, 1))
d1 = diagram_table.cell(2, 0)
set_cell_shading(d1, "DCEEFB")
set_cell_margins(d1, top=25, bottom=25, left=30, right=30)
d1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_border(d1,
    top=(12, "1B2A4A", "single"), bottom=(12, "1B2A4A", "single"),
    left=(12, "1B2A4A", "single"), right=(6, "4A6FA5", "dashed"))
p = d1.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "D-Region (D1)", "Arial", 9, bold=True, color=NAVY)
p2 = d1.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p2, before=1, after=0, line_spacing=1.0)
add_run(p2, "비선형 응력", "Malgun Gothic", 7.5, italic=True, color=STEEL_BLUE)

# B
b = diagram_table.cell(2, 2)
set_cell_shading(b, SECTION_BG)
set_cell_margins(b, top=25, bottom=25, left=30, right=30)
b.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_border(b,
    top=(12, "1B2A4A", "single"), bottom=(12, "1B2A4A", "single"),
    left=(6, "4A6FA5", "dashed"), right=(6, "4A6FA5", "dashed"))
p = b.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "B-Region", "Arial", 9, bold=True, color=DARK_GRAY)
p2 = b.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p2, before=1, after=0, line_spacing=1.0)
add_run(p2, "선형 응력 (평면유지)", "Malgun Gothic", 7.5, italic=True, color=MEDIUM_GRAY)

# D2
diagram_table.cell(2, 3).merge(diagram_table.cell(2, 4))
d2 = diagram_table.cell(2, 3)
set_cell_shading(d2, "DCEEFB")
set_cell_margins(d2, top=25, bottom=25, left=30, right=30)
d2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_border(d2,
    top=(12, "1B2A4A", "single"), bottom=(12, "1B2A4A", "single"),
    left=(6, "4A6FA5", "dashed"), right=(12, "1B2A4A", "single"))
p = d2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "D-Region (D2)", "Arial", 9, bold=True, color=NAVY)
p2 = d2.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p2, before=1, after=0, line_spacing=1.0)
add_run(p2, "비선형 응력", "Malgun Gothic", 7.5, italic=True, color=STEEL_BLUE)

# Row 3: Supports
set_row_height(diagram_table.rows[3], 0.45, "atLeast")
for i in range(5):
    c = diagram_table.cell(3, i)
    set_cell_margins(c, top=0, bottom=0, left=10, right=10)
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

sl = diagram_table.cell(3, 0)
sl.text = ""
p = sl.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "▲", "Arial", 14, bold=True, color=DARK_NAVY)

sr = diagram_table.cell(3, 4)
sr.text = ""
p = sr.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "▲", "Arial", 14, bold=True, color=DARK_NAVY)

# Row 4: Dimension labels
set_row_height(diagram_table.rows[4], 0.5, "atLeast")
for i in range(5):
    c = diagram_table.cell(4, i)
    set_cell_margins(c, top=0, bottom=0, left=10, right=10)
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

ll = diagram_table.cell(4, 0)
ll.text = ""
p = ll.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "지점 R\u2081", "Malgun Gothic", 8, color=MEDIUM_GRAY)

diagram_table.cell(4, 1).merge(diagram_table.cell(4, 3))
dl = diagram_table.cell(4, 1)
dl.text = ""
p = dl.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "|\u2190 h \u2192|\u2190 B-Region \u2192|\u2190 h \u2192|", "Arial", 8, color=STEEL_BLUE)

rl = diagram_table.cell(4, 4)
rl.text = ""
p = rl.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "지점 R\u2082", "Malgun Gothic", 8, color=MEDIUM_GRAY)

# Dimension footnote
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_paragraph_spacing(p, before=2, after=2, line_spacing=1.0)
add_run(p, "\u203B h = \uBD80\uC7AC \uAE4A\uC774(Depth)  |  D = Disturbed  |  B = Bernoulli(\uAC74\uC804\uAD6C\uC5ED)", "Malgun Gothic", 8, italic=True, color=MEDIUM_GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: STM 설계 메커니즘
# Column layout: Component(2.5cm) | Role(4.5cm) | Design(5.5cm) | Node(3.65cm)
# ═══════════════════════════════════════════════════════════════════════════════

add_spacer(doc, 5)
add_section_header(doc, 4, "STM 설계 메커니즘", "Strut-and-Tie Model / KDS 14 20:2024")

TABLE4_W = TABLE2_W  # same as section 2
COL4_WIDTHS = [2.5, 4.5, 5.5, TABLE4_W - 2.5 - 4.5 - 5.5]

stm_table = doc.add_table(rows=4, cols=4)
stm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(stm_table, TABLE4_W)
set_table_col_widths(stm_table, COL4_WIDTHS)

# Header
stm_headers = ["구성요소", "역  할", "설계 핵심", "Node 유형"]
for i, h in enumerate(stm_headers):
    cell = stm_table.cell(0, i)
    set_cell_shading(cell, HEADER_BG)
    set_cell_margins(cell, top=45, bottom=45, left=50, right=50)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, "Malgun Gothic", 9, bold=True, color=WHITE)

stm_data = [
    ["Strut\n(압축대)", "콘크리트 압축 응력장\n하중\u2192절점 전달", "유효 압축 강도 검토\n(Bottle/Prismatic형)", "\u2014"],
    ["Tie\n(인장재)", "철근의 인장 응력 수용\n및 정착", "항복 강도 확보\n정착 길이 검토", "\u2014"],
    ["Node\n(절점)", "Strut+Tie 교차부\n하중 평형 구역", "지압 강도 검토\n정착 상세 확보", "CCC, CCT,\nCTT"],
]

for row_idx, row_data in enumerate(stm_data):
    bg = LIGHT_GRAY_BG if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, text in enumerate(row_data):
        cell = stm_table.cell(row_idx + 1, col_idx)
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=35, bottom=35, left=50, right=50)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        if col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, "Malgun Gothic", 9, bold=True, color=NAVY)
        elif col_idx == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, "Arial", 9, bold=True, color=STEEL_BLUE)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, "Malgun Gothic", 8.5, color=DARK_GRAY)

apply_table_borders(stm_table, HEADER_BG)

# Design core callout
p = doc.add_paragraph()
set_paragraph_spacing(p, before=4, after=2, line_spacing=1.3)
p.paragraph_format.left_indent = Cm(0.4)
add_run(p, "\u25B6 ", "Arial", 9, bold=True, color=ACCENT_GOLD)
add_run(p, "설계 핵심: ", "Malgun Gothic", 9, bold=True, color=NAVY)
add_run(p, "하중 경로(Load Path) 시각화 및 절점부 지압/정착 상세 검토 (KDS 14 20:2024 준수)", "Malgun Gothic", 9, color=DARK_GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: 기술사적 소견
# ═══════════════════════════════════════════════════════════════════════════════

add_spacer(doc, 5)
add_section_header(doc, 5, "기술사적 소견", "Professional Insight")

insight_tbl = doc.add_table(rows=1, cols=1)
insight_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(insight_tbl, TABLE2_W)
cell = insight_tbl.cell(0, 0)
set_cell_shading(cell, "F0F4FF")
set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

insights = [
    ("실무 경험 (17년):", "D-Region 설계 실패는 대개 '철근의 정착' 및 '절점부의 압축 압쇄'에서 기인함."),
    ("설계 착안점:", "전단 설계 시 단순 Vc+Vs 식에 의존하지 말고, STM을 통한 하중의 물리적 거동 파악이 기술사의 핵심 역량임."),
    ("최신 경향:", "2024 KDS 기준은 복잡한 불연속 구간에서 STM 설계를 원칙으로 하며, 특히 절점 상세 설계(Detailing)를 강조함."),
]

for i, (title, desc) in enumerate(insights):
    p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
    set_paragraph_spacing(p, before=2, after=4, line_spacing=1.35)
    add_run(p, f"  \u25A0 {title} ", "Malgun Gothic", 9, bold=True, color=NAVY)
    add_run(p, desc, "Malgun Gothic", 9, color=DARK_GRAY)

# Left-accent border
tbl_el4 = insight_tbl._tbl
tblPr_el4 = tbl_el4.tblPr
brd4 = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="4A6FA5"/>'
    f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{HEADER_BG}"/>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="4A6FA5"/>'
    f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="4A6FA5"/>'
    f'</w:tblBorders>'
)
for existing in tblPr_el4.findall(qn('w:tblBorders')):
    tblPr_el4.remove(existing)
tblPr_el4.append(brd4)


# ═══════════════════════════════════════════════════════════════════════════════
# DIVIDER: 이하 빈칸
# ═══════════════════════════════════════════════════════════════════════════════

add_spacer(doc, 7)
add_horizontal_line(doc, "D0D5E0", 10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=4, after=4)
add_run(p, "\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500  \uc774  \ud558  \ube48  \uce78  \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500", "Malgun Gothic", 10, color=MEDIUM_GRAY)

add_horizontal_line(doc, "D0D5E0", 10)


# ═══════════════════════════════════════════════════════════════════════════════
# STRATEGY SECTION: 문제 분석 및 핵심 전략
# Column layout: Label(2.2cm) | Content(rest) — shrink label, expand content
# ═══════════════════════════════════════════════════════════════════════════════

add_spacer(doc, 7)

# Gold header bar
strat_header = doc.add_table(rows=1, cols=1)
strat_header.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(strat_header, CONTENT_WIDTH_CM)
cell = strat_header.cell(0, 0)
set_cell_shading(cell, GOLD_HEX)
set_cell_margins(cell, top=45, bottom=45, left=100, right=100)
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "\u25A0  \ubb38\uc81c \ubd84\uc11d \ubc0f \ud575\uc2ec \uc804\ub7b5  \u25A0", "Malgun Gothic", 11, bold=True, color=WHITE)
remove_table_borders(strat_header)

# Strategy content — 2 columns with controlled widths
STRAT_LABEL_W = 2.2
STRAT_CONTENT_W = CONTENT_WIDTH_CM - STRAT_LABEL_W

strat_tbl = doc.add_table(rows=4, cols=2)
strat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width_fixed(strat_tbl, CONTENT_WIDTH_CM)
set_table_col_widths(strat_tbl, [STRAT_LABEL_W, STRAT_CONTENT_W])

strategy_items = [
    ("핵심 키워드", "Saint-Venant's Principle  |  STM (Strut-and-Tie Model)  |  Bernoulli Hypothesis  |  h (Depth)"),
    ("반드시 포함", "\u2460 평면유지 가설 파괴 개념  \u2461 D-Region 범위 설정 기준 (부재깊이 h)  \u2462 STM 구성요소 및 설계 핵심"),
    ("유사 기출", "\u2022 108회 1-5 (개구부 설계)  \u2022 125회 2-3 (댑트 보 계산)  \u2022 132회 1-8 (지압 설계)"),
    ("합격 포인트", "단순 암기 답안이 아닌, '\ud558\uc911 \uacbd\ub85c(Load Path)'\uc758 이해를 보여주는 삽도 배치가 당락을 결정함."),
]

for idx, (label, content) in enumerate(strategy_items):
    bg = STRATEGY_BG if idx % 2 == 0 else "FFFFFF"

    # Label cell — compact, gold text
    lc = strat_tbl.cell(idx, 0)
    set_cell_shading(lc, "FFF3E0")
    set_cell_margins(lc, top=45, bottom=45, left=50, right=50)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = lc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, label, "Malgun Gothic", 8.5, bold=True, color=ACCENT_GOLD)

    # Content cell — spacious
    cc = strat_tbl.cell(idx, 1)
    set_cell_shading(cc, bg)
    set_cell_margins(cc, top=45, bottom=45, left=80, right=60)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cc.paragraphs[0]
    add_run(p, content, "Malgun Gothic", 9, color=DARK_GRAY)

apply_table_borders(strat_tbl, GOLD_HEX, inner_h_color="F5DEB3", inner_v_color="F5DEB3")


# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════════

add_spacer(doc, 10)
add_horizontal_line(doc, HEADER_BG, 6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=4, after=0)
add_run(p, "BOSS  |  \uc81c120\ud68c \ud1a0\ubaa9\uad6c\uc870\uae30\uc220\uc0ac  |  1\uad50\uc2dc 1\ubc88  |  D-Region & STM", "Malgun Gothic", 8, color=MEDIUM_GRAY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p2, before=2, after=0)
add_run(p2, "Generated by flowithOS  \u00b7  Professional PE Exam Sub-note System v2", "Arial", 7, italic=True, color=RGBColor(0xAA, 0xAA, 0xAA))


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_FILE)
print(f"\u2705 Document generated successfully: {OUTPUT_FILE}")
print(f"   File size: {os.path.getsize(OUTPUT_FILE):,} bytes")
