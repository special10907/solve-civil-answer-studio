"""
Shared v2 layout framework for PE Exam Sub-note generation.
Refactored: parameterized exam_no/period, literal Korean strings,
            add_flow_diagram(), spacer_before in add_section_header.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Constants ───────────────────────────────────────────────────────────────
CONTENT_WIDTH_CM = 17.0
TABLE_W = CONTENT_WIDTH_CM * 0.95  # ~16.15cm

# Colors
NAVY       = RGBColor(0x1B, 0x2A, 0x4A)
DARK_NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
STEEL_BLUE = RGBColor(0x4A, 0x6F, 0xA5)
ACCENT_RED = RGBColor(0xC0, 0x39, 0x2B)
ACCENT_GOLD= RGBColor(0xD4, 0x8B, 0x0B)
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY= RGBColor(0x66, 0x66, 0x66)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

LIGHT_GRAY_BG = "F2F4F8"
HEADER_BG     = "1B2A4A"
SUB_HEADER_BG = "E8ECF2"
SECTION_BG    = "F7F9FC"
STRATEGY_BG   = "FFF8E1"
GOLD_HEX      = "D48B0B"

BADGE_WIDTH   = 0.7
TITLE_WIDTH   = CONTENT_WIDTH_CM - BADGE_WIDTH
STRAT_LABEL_W = 2.2
STRAT_CONTENT_W = CONTENT_WIDTH_CM - STRAT_LABEL_W


# ─── Low-level utilities ─────────────────────────────────────────────────────

def cm_to_twips(cm):
    return int(cm * 567)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color, style) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{style}" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = parse_xml(f'<w:{side} {nsdecls("w")} w:w="{val}" w:type="dxa"/>')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tw = cm_to_twips(width_cm)
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{tw}" w:type="dxa"/>')
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcPr.append(tcW)

def set_table_col_widths(table, widths_cm):
    tbl = table._tbl
    tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
    for w in widths_cm:
        gridCol = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{cm_to_twips(w)}"/>')
        tblGrid.append(gridCol)
    for existing in tbl.findall(qn('w:tblGrid')):
        tbl.remove(existing)
    tbl.insert(tbl.index(tbl.tblPr) + 1, tblGrid)
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                set_cell_width(row.cells[i], w)

def add_run(paragraph, text, font_name="Malgun Gothic", size=10, bold=False,
            italic=False, color=DARK_GRAY, underline=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if underline:
        run.font.underline = True
    return run

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing

def add_horizontal_line(doc, color="1B2A4A", thickness=6):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, color)
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{thickness}" w:hRule="exact"/>')
    trPr.append(trH)
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)
    return tbl

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left   w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right  w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>')
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(borders)

def set_table_width_fixed(table, width_cm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tw = cm_to_twips(width_cm)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{tw}" w:type="dxa"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblPr.append(tblLayout)

def set_row_height(row, height_cm, rule="exact"):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{cm_to_twips(height_cm)}" w:hRule="{rule}"/>')
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)
    trPr.append(trH)

def apply_table_borders(table, outer_color, inner_h_color="D0D5E0", inner_v_color="D0D5E0",
                         outer_sz=8, inner_sz=4):
    tbl_el = table._tbl
    tblPr_el = tbl_el.tblPr
    brd = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:left   w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:bottom w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:right  w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:insideH w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_h_color}"/>'
        f'  <w:insideV w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_v_color}"/>'
        f'</w:tblBorders>')
    for existing in tblPr_el.findall(qn('w:tblBorders')):
        tblPr_el.remove(existing)
    tblPr_el.append(brd)

def add_spacer(doc, pts=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=pts, after=0)
    return p


# ─── High-level components ───────────────────────────────────────────────────

def create_document():
    """Create a new document with A4 page setup and default styles."""
    doc = Document()
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.top_margin   = Cm(1.5)
    section.bottom_margin= Cm(1.5)
    section.left_margin  = Cm(2.0)
    section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    style.paragraph_format.space_before  = Pt(0)
    style.paragraph_format.space_after   = Pt(0)
    style.paragraph_format.line_spacing  = 1.15
    return doc


def add_exam_header(doc, question_number, question_title,
                    exam_no=120, period=1, examinee_name="B  O  S  S"):
    """
    Add the standard 3-row exam header block.

    Parameters
    ----------
    exam_no       : 회차 번호 (default 120)
    period        : 교시 번호 (default 1)
    examinee_name : 수험자 성명 표시 문자열
    """
    header_table = doc.add_table(rows=3, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_fixed(header_table, CONTENT_WIDTH_CM)
    set_table_col_widths(header_table, [7.0, 5.0, 5.0])

    # Row 0: Title bar
    row0 = header_table.rows[0]
    set_row_height(row0, 1.1)
    for i in range(3):
        cell = row0.cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell, top=50, bottom=50, left=100, right=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row0.cells[0].merge(row0.cells[2])
    merged = row0.cells[0]
    merged.text = ""
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"[ 제 {exam_no} 회 ]   ", "Arial", 13, bold=True, color=WHITE)
    add_run(p, "토  목  구  조  기  술  사", "Malgun Gothic", 14, bold=True, color=WHITE)
    add_run(p, f"   ( 제 {period} 교 시 )", "Arial", 13, bold=True, color=WHITE)

    # Row 1: Info bar
    row1 = header_table.rows[1]
    set_row_height(row1, 0.8)
    for i in range(3):
        cell = row1.cells[i]
        set_cell_shading(cell, SUB_HEADER_BG)
        set_cell_margins(cell, top=35, bottom=35, left=100, right=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    c = row1.cells[0]; c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "성 명 :  ", "Malgun Gothic", 10, bold=True, color=NAVY)
    add_run(p, examinee_name, "Arial", 11, bold=True, color=ACCENT_RED)

    c = row1.cells[1]; c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "문제번호 :  ", "Malgun Gothic", 10, bold=True, color=NAVY)
    add_run(p, str(question_number), "Arial", 12, bold=True, color=ACCENT_RED)

    c = row1.cells[2]; c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "Page  ", "Arial", 9, color=MEDIUM_GRAY)
    add_run(p, "1 / 1", "Arial", 9, bold=True, color=NAVY)

    # Row 2: Question title
    row2 = header_table.rows[2]
    set_row_height(row2, 0.9)
    for i in range(3):
        cell = row2.cells[i]
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row2.cells[0].merge(row2.cells[2])
    merged = row2.cells[0]
    merged.text = ""
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f"문 {question_number}.  ", "Malgun Gothic", 12, bold=True, color=DARK_NAVY)
    add_run(p, question_title,             "Malgun Gothic", 12, bold=True, color=DARK_NAVY)

    apply_table_borders(header_table, HEADER_BG, inner_h_color="D0D5E0", inner_v_color="D0D5E0",
                         outer_sz=12, inner_sz=4)
    tbl_pr = header_table._tbl.tblPr
    for brd in tbl_pr.findall(qn('w:tblBorders')):
        for iv in brd.findall(qn('w:insideV')):
            iv.set(qn('w:val'), 'none')
    add_spacer(doc, 5)


def add_section_header(doc, number, title_kr, title_en, spacer_before=5):
    """
    Compact numbered section header with underline.

    Parameters
    ----------
    spacer_before : 섹션 앞 spacer 크기(pt). 0이면 spacer 없음.
    """
    if spacer_before > 0:
        add_spacer(doc, spacer_before)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width_fixed(tbl, CONTENT_WIDTH_CM)
    set_table_col_widths(tbl, [BADGE_WIDTH, TITLE_WIDTH])
    remove_table_borders(tbl)
    badge = tbl.cell(0, 0)
    set_cell_shading(badge, HEADER_BG)
    badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(badge, top=35, bottom=35, left=0, right=0)
    p = badge.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(number), "Arial", 11, bold=True, color=WHITE)
    title = tbl.cell(0, 1)
    set_cell_margins(title, top=35, bottom=35, left=100, right=40)
    title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(title, bottom=(8, HEADER_BG, "single"))
    p = title.paragraphs[0]
    add_run(p, f"{title_kr} ", "Malgun Gothic", 11, bold=True, color=DARK_NAVY)
    add_run(p, f"({title_en})", "Arial", 9, italic=True, color=STEEL_BLUE)
    return tbl


def add_content_line(doc, number, text, bold_prefix="", indent_level=0):
    """Numbered content line with circled number marker."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line_spacing=1.3)
    p.paragraph_format.left_indent = Cm(0.4 + (indent_level * 0.4))
    markers = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩"]
    marker = markers[min(number - 1, len(markers) - 1)] if number > 0 else "•"
    add_run(p, f" {marker} ", "Arial", 9, color=STEEL_BLUE)
    if bold_prefix:
        add_run(p, bold_prefix, "Malgun Gothic", 10, bold=True, color=DARK_NAVY)
    add_run(p, text, "Malgun Gothic", 10, color=DARK_GRAY)
    return p


def add_content_sub(doc, text, bold_prefix="", indent=1.0):
    """Sub-content paragraph with description text."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.3)
    p.paragraph_format.left_indent = Cm(indent)
    if bold_prefix:
        add_run(p, bold_prefix, "Malgun Gothic", 10, bold=True, color=DARK_NAVY)
    add_run(p, text, "Malgun Gothic", 10, color=DARK_GRAY)
    return p


def add_content_sub_highlight(doc, prefix_text, highlight_text, suffix=".", indent=1.0):
    """Sub-content with a red-highlighted keyword phrase."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.3)
    p.paragraph_format.left_indent = Cm(indent)
    add_run(p, prefix_text,    "Malgun Gothic", 10, color=DARK_GRAY)
    add_run(p, highlight_text, "Malgun Gothic", 10, bold=True, color=ACCENT_RED)
    add_run(p, suffix,         "Malgun Gothic", 10, color=DARK_GRAY)
    return p


def add_data_table(doc, headers, data, col_widths, first_col_center=True):
    """Create a professional data table with navy header and alternating rows."""
    total_w  = sum(col_widths)
    num_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(data), cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_fixed(tbl, total_w)
    set_table_col_widths(tbl, col_widths)

    # Header
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell, top=45, bottom=45, left=60, right=60)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, "Malgun Gothic", 9, bold=True, color=WHITE)

    # Data rows (alternating bg)
    for row_idx, row_data in enumerate(data):
        bg = LIGHT_GRAY_BG if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = tbl.cell(row_idx + 1, col_idx)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            if col_idx == 0 and first_col_center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, text, "Malgun Gothic", 9, bold=True, color=NAVY)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_run(p, text, "Malgun Gothic", 9, color=DARK_GRAY)

    apply_table_borders(tbl, HEADER_BG)
    return tbl


def add_flow_diagram(doc, steps, width_ratio=0.9):
    """
    Add a horizontal flow diagram using a table of step boxes with arrows.

    Parameters
    ----------
    steps : list of (label, bg_color_hex, RGBColor) tuples for each box.
            Arrow cells are inserted automatically between boxes.
    width_ratio : fraction of CONTENT_WIDTH_CM to use (default 0.9).

    Example
    -------
    add_flow_diagram(doc, [
        ("Cl⁻/H₂O\n침투",  "FFEBEE", ACCENT_RED),
        ("Pit 발생 +\nH⁰ 축적", "FFF3E0", ACCENT_GOLD),
        ("SCC/HE\n파단",   "DCEEFB", NAVY),
    ])
    """
    # Interleave arrows between boxes
    cells_data = []
    for i, step in enumerate(steps):
        cells_data.append(step)
        if i < len(steps) - 1:
            cells_data.append(("→", SECTION_BG, MEDIUM_GRAY))

    n = len(cells_data)
    total_w = CONTENT_WIDTH_CM * width_ratio
    col_w = total_w / n

    diag = doc.add_table(rows=1, cols=n)
    diag.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_fixed(diag, total_w)
    set_table_col_widths(diag, [col_w] * n)
    remove_table_borders(diag)

    for i, (txt, bg, clr) in enumerate(cells_data):
        c = diag.cell(0, i)
        set_cell_shading(c, bg)
        set_cell_margins(c, top=35, bottom=35, left=20, right=20)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        is_arrow = (txt == "→")
        add_run(p, txt,
                "Arial" if is_arrow else "Malgun Gothic",
                16 if is_arrow else 9,
                bold=True, color=clr)
    return diag


def add_insight_box(doc, insights):
    """Add left-accent bordered insight box with bullet items."""
    insight_tbl = doc.add_table(rows=1, cols=1)
    insight_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_fixed(insight_tbl, TABLE_W)
    cell = insight_tbl.cell(0, 0)
    set_cell_shading(cell, "F0F4FF")
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    for i, (title, desc) in enumerate(insights):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        set_paragraph_spacing(p, before=2, after=4, line_spacing=1.35)
        add_run(p, f"  ■ {title} ", "Malgun Gothic", 9, bold=True, color=NAVY)
        add_run(p, desc, "Malgun Gothic", 9, color=DARK_GRAY)

    tbl_el = insight_tbl._tbl
    tblPr_el = tbl_el.tblPr
    brd = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="6"  w:space="0" w:color="4A6FA5"/>'
        f'  <w:left   w:val="single" w:sz="18" w:space="0" w:color="{HEADER_BG}"/>'
        f'  <w:bottom w:val="single" w:sz="6"  w:space="0" w:color="4A6FA5"/>'
        f'  <w:right  w:val="single" w:sz="6"  w:space="0" w:color="4A6FA5"/>'
        f'</w:tblBorders>')
    for existing in tblPr_el.findall(qn('w:tblBorders')):
        tblPr_el.remove(existing)
    tblPr_el.append(brd)
    return insight_tbl


def add_divider(doc):
    """Add the '이하 빈칸' divider section."""
    add_spacer(doc, 7)
    add_horizontal_line(doc, "D0D5E0", 10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=4)
    add_run(p, "─" * 14 + "  이  하  빈  칸  " + "─" * 14,
            "Malgun Gothic", 10, color=MEDIUM_GRAY)
    add_horizontal_line(doc, "D0D5E0", 10)


def add_strategy_section(doc, items):
    """
    Add the gold strategy analysis section.
    items: list of (label, content) tuples.
    """
    add_spacer(doc, 7)
    # Gold header bar
    strat_header = doc.add_table(rows=1, cols=1)
    strat_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_fixed(strat_header, CONTENT_WIDTH_CM)
    cell = strat_header.cell(0, 0)
    set_cell_shading(cell, GOLD_HEX)
    set_cell_margins(cell, top=45, bottom=45, left=100, right=100)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "■  문제 분석 및 핵심 전략  ■",
            "Malgun Gothic", 11, bold=True, color=WHITE)
    remove_table_borders(strat_header)

    # Strategy table
    strat_tbl = doc.add_table(rows=len(items), cols=2)
    strat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_fixed(strat_tbl, CONTENT_WIDTH_CM)
    set_table_col_widths(strat_tbl, [STRAT_LABEL_W, STRAT_CONTENT_W])

    for idx, (label, content) in enumerate(items):
        bg = STRATEGY_BG if idx % 2 == 0 else "FFFFFF"
        lc = strat_tbl.cell(idx, 0)
        set_cell_shading(lc, "FFF3E0")
        set_cell_margins(lc, top=45, bottom=45, left=50, right=50)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = lc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label, "Malgun Gothic", 8.5, bold=True, color=ACCENT_GOLD)

        cc = strat_tbl.cell(idx, 1)
        set_cell_shading(cc, bg)
        set_cell_margins(cc, top=45, bottom=45, left=80, right=60)
        cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cc.paragraphs[0]
        add_run(p, content, "Malgun Gothic", 9, color=DARK_GRAY)

    apply_table_borders(strat_tbl, GOLD_HEX, inner_h_color="F5DEB3", inner_v_color="F5DEB3")


def add_footer(doc, question_number, short_title, exam_no=120, period=1):
    """Add footer with branding."""
    add_spacer(doc, 10)
    add_horizontal_line(doc, HEADER_BG, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=0)
    add_run(p, f"BOSS  |  제{exam_no}회 토목구조기술사  |  {period}교시 {question_number}번  |  {short_title}",
            "Malgun Gothic", 8, color=MEDIUM_GRAY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, before=2, after=0)
    add_run(p2, "Generated by flowithOS  ·  Professional PE Exam Sub-note System v2",
            "Arial", 7, italic=True, color=RGBColor(0xAA, 0xAA, 0xAA))


def add_callout_box(doc, title, body_text, bg_color="FFF3E0", border_color="FFB74D",
                    title_color=None):
    """Add a colored callout/instruction box."""
    if title_color is None:
        title_color = RGBColor(0xE6, 0x5C, 0x00)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_fixed(tbl, TABLE_W)
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, bg_color)
    set_cell_margins(cell, top=50, bottom=50, left=100, right=100)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.2)
    add_run(p, f"  {title}  ", "Malgun Gothic", 9, bold=True, color=title_color)
    if body_text:
        p2 = cell.add_paragraph()
        set_paragraph_spacing(p2, before=2, after=0, line_spacing=1.25)
        add_run(p2, f"  {body_text}", "Malgun Gothic", 8, color=DARK_GRAY)

    tbl_el = tbl._tbl
    tblPr_el = tbl_el.tblPr
    brd = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'  <w:left   w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'  <w:right  w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'</w:tblBorders>')
    for existing in tblPr_el.findall(qn('w:tblBorders')):
        tblPr_el.remove(existing)
    tblPr_el.append(brd)
    return tbl


def save_document(doc, output_path):
    """Save and report file size."""
    doc.save(output_path)
    size = os.path.getsize(output_path)
    print(f"✅ Generated: {os.path.basename(output_path)} ({size:,} bytes)")
